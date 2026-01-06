import tkinter as tk
from tkinter import font, filedialog, messagebox, ttk
from tkinter import scrolledtext  # <--- Import para a Ajuda e o Preview
import os
import sys
import subprocess
import json
import threading
import editortkinter
from docx import Document         # <--- Import para ler o Word no Preview

# --- TEXTO DO README EMBUTIDO ---
README_TEXT = """
# AutoFoto Report - Organizador e Importador de Fotos para Relatórios

**Versão 4.2.3** | Criado por: Pedro H.G.C Vidal

## Visão Geral

O AutoFoto Report é uma ferramenta projetada para eliminar o trabalho manual e repetitivo de adicionar, editar e formatar fotos em relatórios do Microsoft Word. A solução automatiza o fluxo de trabalho, desde a organização das imagens até a inserção no documento, incluindo um editor de imagens integrado para garantir a qualidade, ordem e padronização.

## Principais Funcionalidades

* **Editor de Imagem Integrado (Mestre):** Uma interface gráfica única permite visualizar todas as fotos em fila. Você pode cortar, girar, ajustar o brilho e reordenar cada imagem de forma interativa.
* **Automação de Layout:** Insere as fotos processadas no local exato do documento Word (selecionado visualmente no launcher), com formatação de tamanho, alinhamento e espaçamento pré-configurada.
* **Modos Flexíveis de Organização:**
    * **Modo Lotes:** Processa fotos organizadas em subpastas. Insere títulos automaticamente para cada lote no documento.
    * **Modo Pasta Única:** Processa todas as fotos encontradas diretamente dentro da pasta selecionada, tratando-as como um único lote.
* **Seleção Visual do Local:** Permite ao usuário abrir uma pré-visualização do documento Word e clicar no parágrafo exato onde as fotos devem ser inseridas, eliminando a necessidade de textos-âncora.

## Fluxo de Trabalho Passo a Passo

**Passo 1: Preparação (Apenas 1ª vez)**

A única preparação necessária é ter um documento Word (`.docx`) e uma pasta com as fotos que você deseja inserir. As fotos podem estar organizadas em subpastas (Modo Lotes) ou todas juntas na mesma pasta (Modo Pasta Única).

**Passo 2: Execução do AutoFoto Report**

1.  Execute o programa `AutoFotoReport.exe`.
2.  Na janela principal:
    * **`1. Selecionar Relatório Word...`**: Escolha o `.docx` base.
    * **`2. Selecionar Pasta de Fotos...`**: Escolha a pasta principal que contém suas fotos (seja a que tem os lotes dentro, ou a que tem as fotos diretamente).
    * **`3. Definir Local de Saída...`**: Escolha onde salvar o novo relatório (ex: `Relatorio_Final_COM_FOTOS.docx`).
    * **`4. Escolher Local de Inserção no Word...`**: (Este botão só ativa após o Passo 1). Clique nele para abrir uma janela de pré-visualização do seu Word. Role e **clique no parágrafo** após o qual as fotos devem ser inseridas. Confirme a seleção.
3.  Clique no botão verde **"Iniciar Importação e Edição de Fotos"**.
4.  Uma caixa de diálogo perguntará se você deseja editar as imagens.
    * **`Sim`**: A janela do "Editor de Fotos Mestre" será aberta.
    * **`Não`**: O script pulará a edição e inserirá as fotos originais.

## Guia Rápido do Editor Mestre

O Editor Mestre é dividido em três seções:

* **1. Lista de Fotos (Esquerda)**
    * **Lista:** Mostra a ordem atual das fotos (`Lote / Arquivo`). Clique em qualquer foto para carregá-la no editor.
    * **`Mover Cima ↑` / `Mover Baixo ↓`:** Move a foto selecionada uma posição para cima ou para baixo. **Só funciona dentro do mesmo lote.**
    * **`Reordenar Lote (1,2,3...)`:** Abre uma nova janela *apenas* para o lote da foto selecionada. Nela, você pode:
        1. Clicar no nome de uma foto para pré-visualizá-la à direita.
        2. Digitar a nova ordem numérica (1, 2, 3...) nos campos à esquerda.
        3. Clicar em "Confirmar Ordem" para aplicar a reordenação em massa a esse lote.

* **2. Editor (Centro)**
    * **Canvas:** Exibe a imagem selecionada. Clique e arraste para desenhar o retângulo verde de corte.
    * **Overlay Azul:** Mostra o tamanho e posição do último corte que você aplicou.
    * **Slider `Luminosidade`:** Ajusta o brilho da imagem atual.
    * **`Reset`:** Retorna o brilho para o padrão (100%).

* **3. Botões de Ação (Inferior)**
    * **`Finalizar Edição`:** Salva todas as alterações e fecha o editor para gerar o documento.
    * **`Girar 90° Esq / Dir`:** Rotaciona a imagem atual. A rotação será aplicada no documento final.
    * **`Aplicar Corte e Próximo >>`:** Salva o corte atual (retângulo verde) e avança para a próxima foto.
    * **`Pular (Usar Original) e Próximo >>`:** Ignora o corte para a foto atual (mas mantém a rotação/brilho, se aplicados) e avança.
    * **`Aplicar Último Corte`:** Aplica as dimensões do overlay azul ao corte atual. Não avança.
    * **`Replicar para Demais`:** Aplica o **corte**, **rotação** e **brilho** da foto atual para todas as fotos seguintes na lista e finaliza a edição.

## Passo 3: Resultado Final

Após a conclusão do processo, o seu novo documento Word (com todas as fotos inseridas, editadas e formatadas na ordem correta) será salvo no local que você definiu. O programa tentará abrir o arquivo automaticamente.
"""
# --- FIM DO TEXTO DO README ---


# --- Lógica de Caminho e Nomes de Arquivos ---
if getattr(sys, 'frozen', False):
    application_path = os.path.dirname(sys.executable)
else:
    application_path = os.path.dirname(os.path.abspath(__file__))
os.chdir(application_path)

SETTINGS_FILE_NAME = "settings.json"
# (Outros nomes de arquivos não são mais necessários)

# --- Variáveis Globais para o Ponto de Inserção ---
g_config_data = {}
g_all_photos_by_lote = {}
g_selected_paragraph_index = None # Salva o índice do parágrafo clicado

# --- Funções de Ajuda e Configuração ---

def show_help_window():
    """Abre uma janela Toplevel para exibir o README com formatação."""
    help_win = tk.Toplevel(window)
    help_win.title("Ajuda - AutoFoto Report")
    help_win.geometry("700x550")
    help_win.grab_set()
    help_win.resizable(True, True)

    st = scrolledtext.ScrolledText(help_win, wrap=tk.WORD, padx=10, pady=10)
    st.pack(fill=tk.BOTH, expand=True)

    # 1. Tenta criar fontes ideais
    try:
        h1_font = font.Font(family="Helvetica", size=16, weight="bold")
        h2_font = font.Font(family="Helvetica", size=14, weight="bold")
        bold_font = font.Font(family="Helvetica", size=10, weight="bold")
        body_font = font.Font(family="Helvetica", size=10)
    except Exception:
        print("Aviso: Fonte Helvetica não encontrada. Usando fontes padrão.")
        h1_font = font.Font(size=16, weight="bold")
        h2_font = font.Font(size=14, weight="bold")
        bold_font = font.Font(size=10, weight="bold")
        body_font = font.Font(size=10)

    # 2. Configura as tags
    st.tag_configure("h1", font=h1_font, spacing1=5, spacing3=10)
    st.tag_configure("h2", font=h2_font, spacing1=10, spacing3=8)
    st.tag_configure("bold", font=bold_font)
    st.tag_configure("body", font=body_font, spacing3=2)
    st.tag_configure("bullet", font=body_font, lmargin1=20, lmargin2=35, spacing3=2)
        
    # 3. Processar o README_TEXT linha por linha
    st.configure(state='normal')
    lines = README_TEXT.strip().split('\n')
    for line in lines:
        clean_line = line.strip().lstrip()
        tags_to_apply = ["body"]
        text_to_insert = clean_line

        if clean_line.startswith("# "):
            text_to_insert = clean_line[2:]
            tags_to_apply = ["h1"]
        elif clean_line.startswith("## "):
            text_to_insert = clean_line[3:]
            tags_to_apply = ["h2"]
        elif clean_line.startswith("* "):
            text_to_insert = "•\t" + clean_line[2:]
            tags_to_apply = ["bullet"]
        
        st.insert(tk.END, text_to_insert + '\n', tags_to_apply)

    # 4. Aplica o negrito (Pós-processamento)
    start_index = "1.0"
    while True:
        start_pos = st.search("**", start_index, stopindex=tk.END)
        if not start_pos: break
        end_pos = st.search("**", f"{start_pos}+2c", stopindex=tk.END)
        if not end_pos: break
        st.tag_add("bold", f"{start_pos}+2c", end_pos)
        st.delete(end_pos, f"{end_pos}+2c")
        st.delete(start_pos, f"{start_pos}+2c")
        start_index = end_pos

    st.configure(state='disabled')
    close_button = tk.Button(help_win, text="Fechar", command=help_win.destroy)
    close_button.pack(pady=10)

def load_settings():
    # 1. Define os padrões (Defaults) que SEMPRE devem existir
    # (Aqui incluímos todas as opções novas)
    settings = {
        "TITLE_FONT_SIZE_PT": 16, "TITLE_BOLD": True, "CENTER_TITLES": True,
        "CENTER_IMAGES": True, "SPACE_ABOVE_TITLES_PT": 36,
        "SPACE_ABOVE_FIRST_IMAGE_PT": 35, "SPACE_BETWEEN_IMAGES_PT": 24,
        "MAX_IMAGE_WIDTH_CM": 16.0, "MAX_IMAGE_HEIGHT_CM": 22.0,
        "FORCE_RESIZE_SMALLER_IMAGES": False,
        "SIDE_BY_SIDE_LAYOUT": False,
        "ADD_TIMESTAMP": False # <--- Garante que esta chave exista
    }

    # 2. Tenta carregar do arquivo e mesclar (merge) com os padrões
    if os.path.exists(SETTINGS_FILE_NAME):
        try:
            with open(SETTINGS_FILE_NAME, 'r') as f:
                saved_data = json.load(f)
                # O .update() pega o que foi salvo e joga por cima dos padrões.
                # Chaves novas (como ADD_TIMESTAMP) que não estão no arquivo
                # permanecem com o valor padrão do código.
                settings.update(saved_data) 
        except (json.JSONDecodeError, IOError):
            pass # Se o arquivo estiver ruim, usa apenas os padrões

    return settings

def save_settings(new_data):
    """Salva ou atualiza configurações no arquivo JSON sem apagar as existentes."""
    current_settings = {}  
    # 1. Tenta carregar o que já existe
    if os.path.exists(SETTINGS_FILE_NAME):
        try:
            with open(SETTINGS_FILE_NAME, 'r') as f:
                current_settings = json.load(f)
        except:
            pass # Se der erro, começa com dicionário vazio
    # 2. Atualiza com os novos dados
    current_settings.update(new_data)
    # 3. Salva tudo de volta
    with open(SETTINGS_FILE_NAME, 'w') as f:
        json.dump(current_settings, f, indent=4)

def open_settings_window():
    settings = load_settings()
    settings_window = tk.Toplevel(window)
    settings_window.title("Configurações de Layout")
    settings_window.grab_set()
    settings_window.resizable(False, False)
    
    # 1. Define quais configurações são válidas e seus nomes
    friendly_names = {
        "TITLE_FONT_SIZE_PT": "Tam. Fonte Título (pt):", 
        "TITLE_BOLD": "Título em Negrito:",
        "CENTER_TITLES": "Centralizar Títulos:", 
        "CENTER_IMAGES": "Centralizar Imagens:",
        "SPACE_ABOVE_TITLES_PT": "Espaço Acima Título (pt):", 
        "SPACE_ABOVE_FIRST_IMAGE_PT": "Espaço Acima 1ª Imagem (pt):",
        "SPACE_BETWEEN_IMAGES_PT": "Espaço Entre Imagens (pt):", 
        "MAX_IMAGE_WIDTH_CM": "Largura Máx. Imagem (cm):",
        "MAX_IMAGE_HEIGHT_CM": "Altura Máx. Imagem (cm):", 
        "FORCE_RESIZE_SMALLER_IMAGES": "Forçar Redim. Img Pequenas:",
        "SIDE_BY_SIDE_LAYOUT": "Layout Lado a Lado (2 col):",
        "ADD_TIMESTAMP": "Adicionar Carimbo de Data/Hora:"
    }

    # 2. Cria variáveis APENAS para as configurações que estão no friendly_names
    vars = {}
    for key, value in settings.items():
        # --- FILTRO DE SEGURANÇA ---
        if key not in friendly_names:
            continue # Ignora caminhos salvos ou lixo no json
        # ---------------------------

        if isinstance(value, bool):
            vars[key] = tk.BooleanVar(value=value)
        elif isinstance(value, float):
            vars[key] = tk.DoubleVar(value=value)
        else:
            # Tenta converter para int, se falhar (ex: string), ignora
            try:
                int_val = int(value)
                vars[key] = tk.IntVar(value=int_val)
            except (ValueError, TypeError):
                continue

    frame = tk.Frame(settings_window, padx=15, pady=15)
    frame.pack()

    row = 0
    # 3. Gera a interface
    for key, label_text in friendly_names.items():
        if key in vars: # Só mostra se a variável foi criada com sucesso
            var = vars[key]
            tk.Label(frame, text=label_text).grid(row=row, column=0, sticky="w", pady=2)
            
            if isinstance(var, tk.BooleanVar):
                tk.Checkbutton(frame, variable=var).grid(row=row, column=1, sticky="w")
            else:
                tk.Entry(frame, textvariable=var, width=10).grid(row=row, column=1, sticky="w")
            row += 1

    def on_save():
        # Salva apenas o que foi editado, mantendo o resto (caminhos) intacto no arquivo
        updates = {key: var.get() for key, var in vars.items()}
        save_settings(updates)
        settings_window.destroy()

    tk.Button(frame, text="Salvar e Fechar", command=on_save, font=("Helvetica", 10, "bold")).grid(row=row, columnspan=2, pady=10)


# --- Funções de Seleção de Arquivo e Preview ---

def select_input_file():
    filepath = filedialog.askopenfilename(title="Selecione o arquivo Word", filetypes=(("Documentos Word", "*.docx"),))
    if filepath: 
        input_file_var.set(filepath)
        # Habilita o botão de seleção de local
        insertion_btn.config(state=tk.NORMAL)
        # Limpa seleção antiga
        global g_selected_paragraph_index
        g_selected_paragraph_index = None
        insertion_point_var.set("[Nenhum local de inserção selecionado]")

def select_photos_folder():
    folderpath = filedialog.askdirectory(title="Selecione a pasta de fotos (ex: 2_fotos_lotes)")
    if folderpath: photos_folder_var.set(folderpath)

def select_output_file():
    filepath = filedialog.asksaveasfilename(title="Defina o local de saída", filetypes=(("Documentos Word", "*.docx"),), defaultextension=".docx", initialfile="Relatorio_COM_FOTOS.docx")
    if filepath: output_file_var.set(filepath)

def open_document_preview():
    """
    Abre uma janela Toplevel que exibe o conteúdo do docx
    e permite ao usuário clicar para selecionar um ponto de inserção.
    """
    global g_selected_paragraph_index
    
    doc_path = input_file_var.get()
    if not os.path.exists(doc_path):
        messagebox.showerror("Erro", "Arquivo Word não encontrado. Selecione o arquivo no Passo 1 primeiro.")
        return

    preview_win = tk.Toplevel(window)
    preview_win.title("Selecione o Ponto de Inserção")
    preview_win.geometry("800x600")
    preview_win.grab_set()

    tk.Label(preview_win, text="Role pelo documento e clique APÓS o parágrafo onde as fotos devem ser inseridas.", padx=10, pady=10).pack(fill=tk.X)

    st = scrolledtext.ScrolledText(preview_win, wrap=tk.WORD, padx=10, pady=10)
    st.pack(fill=tk.BOTH, expand=True)
    
    # Define uma tag para o item selecionado
    st.tag_configure("highlight", background="yellow", relief="raised")
    
    current_selection_tag = None

    def on_paragraph_click(event, paragraph_index, paragraph_text):
        """Chamado quando um parágrafo é clicado."""
        nonlocal current_selection_tag
        global g_selected_paragraph_index
        
        # Remove o highlight antigo
        if current_selection_tag:
            st.tag_remove("highlight", "1.0", tk.END)
        
        # Adiciona o highlight novo
        tag_name = f"p_{paragraph_index}"
        st.tag_add("highlight", f"{tag_name}.first", f"{tag_name}.last")
        current_selection_tag = tag_name
        
        # Salva a seleção
        g_selected_paragraph_index = paragraph_index
        
        # Atualiza o texto de status na janela principal
        insertion_point_var.set(f"Inserir Após Parágrafo {paragraph_index}: \"{paragraph_text[:50]}...\"")
        
    
    # --- Carrega o Documento ---
    st.configure(state='normal')
    try:
        doc = Document(doc_path)
        if not doc.paragraphs:
            st.insert(tk.END, "(O documento está vazio ou não contém parágrafos legíveis)")
            
        for i, p in enumerate(doc.paragraphs):
            # Não ignora mais linhas em branco, para manter o índice correto
            text = p.text if p.text.strip() else "[...Linha em Branco...]"
            text += "\n\n" # Adiciona espaço extra para legibilidade
            tag_name = f"p_{i}" # Tag única para cada parágrafo
            
            st.insert(tk.END, text, (tag_name,))
            
            # Associa o clique à tag
            st.tag_bind(tag_name, "<Button-1>", 
                        lambda e, index=i, text=p.text: on_paragraph_click(e, index, text if text.strip() else "[Linha em Branco]"))
        
        st.configure(state='disabled')
    except Exception as e:
        st.delete("1.0", tk.END)
        st.insert(tk.END, f"Erro ao ler o documento Word:\n{e}\n\nVerifique se o arquivo não está corrompido ou aberto em outro programa.", "error")

    # Botão para fechar
    def confirm_selection():
        if g_selected_paragraph_index is None:
            messagebox.showwarning("Nenhuma Seleção", "Por favor, clique em um parágrafo no texto para selecioná-lo.")
            return
        preview_win.destroy()
        
    close_button = tk.Button(preview_win, text="Confirmar e Fechar", command=confirm_selection, font=("Helvetica", 10, "bold"))
    close_button.pack(pady=10)


# --- Funções de Controle de Thread ---

def run_importer():
    """
    Função principal do botão "Iniciar". Valida os campos
    e inicia a Fase 1.
    """
    input_doc, photos_folder, output_doc = input_file_var.get(), photos_folder_var.get(), output_file_var.get()
    if "[Nenhum" in input_doc or not input_doc or "[Nenhuma" in photos_folder or not photos_folder or "[Nenhum" in output_doc or not output_doc:
        status_var.set("ERRO: Preencha os campos 1, 2 e 3!")
        return

    # Verifica se o local de inserção foi selecionado
    if g_selected_paragraph_index is None:
        status_var.set("ERRO: Por favor, use o 'Passo 4' para escolher um local de inserção no documento.")
        return
            
    user_wants_to_edit = messagebox.askyesno("Editor", "Deseja cortar/editar as fotos interativamente?")

    layout_settings = load_settings()
    config_data = {
        "input_docx": input_doc, 
        "photos_main_folder": photos_folder, 
        "output_docx": output_doc, 
        "layout": layout_settings,
        "paragraph_index": g_selected_paragraph_index # Passa o índice
    }
    
    start_phase_1_analysis(config_data, user_wants_to_edit)

def start_phase_1_analysis(config_data, editing_enabled):
    """Inicia a thread da Fase 1 (Análise de Pastas)."""
    global g_config_data
    g_config_data = config_data 
    status_var.set("Iniciando Fase 1: Análise de pastas...")
    run_button.config(state=tk.DISABLED)
    progress_bar['mode'] = 'indeterminate' # Garante que está em modo 'indeterminate'
    progress_bar.pack(pady=5)
    progress_bar.start()
    threading.Thread(
        target=_thread_task_fase_1,
        args=(config_data['photos_main_folder'], editing_enabled),
        daemon=True
    ).start()

def _thread_task_fase_1(photos_folder_path, editing_enabled):
    """THREAD WORKER: Executa a Fase 1 (Análise)."""
    try:
        def update_status(message):
            window.after(0, status_var.set, message)
        all_photos, flat_list = editortkinter.fase_1_analise(photos_folder_path, update_status)
        global g_all_photos_by_lote
        g_all_photos_by_lote = all_photos
        window.after(0, start_phase_2_editor, flat_list, editing_enabled)
    except Exception as e:
        window.after(0, on_process_error, e)

def start_phase_2_editor(flat_photo_list, editing_enabled):
    """GUI THREAD: Abre o editor ou pula para a Fase 3."""
    progress_bar.stop()
    if editing_enabled:
        status_var.set("Aguardando edição de fotos...")
        editor = editortkinter.MasterCropEditor(
            window, 
            flat_photo_list,
            g_all_photos_by_lote,
            start_phase_3_generation # Função a chamar ao fechar
        )
    else:
        status_var.set("Edição desativada, pulando para a geração...")
        decisions = {path: {"crop": "skip", "rotation": 0, "brightness": 1.0} for path in flat_photo_list}
        # Passa a lista original e as decisões 'skip'
        start_phase_3_generation(flat_photo_list, decisions)
        
def start_phase_3_generation(reordered_flat_list, decisions):
    """Inicia a thread da Fase 3 (Geração do DOCX)."""
    status_var.set("Reconstruindo ordem dos lotes...")
    
    reordered_photos_by_lote = {}
    for path in reordered_flat_list:
        lote_name = os.path.basename(os.path.dirname(path))
        if lote_name not in reordered_photos_by_lote:
            reordered_photos_by_lote[lote_name] = []
        reordered_photos_by_lote[lote_name].append(path)
        
    status_var.set("Iniciando Fase 3: Geração do documento...")
    
    # Configura a barra para modo 'determinate'
    progress_bar['mode'] = 'determinate'
    progress_bar['value'] = 0 
    progress_bar.pack(pady=5) # Garante que está visível
    
    threading.Thread(
        target=_thread_task_fase_3,
        args=(g_config_data, reordered_photos_by_lote, decisions), 
        daemon=True
    ).start()

def _thread_task_fase_3(config_data, all_photos, decisions):
    """THREAD WORKER: Executa a Fase 3 (Geração)."""
    try:
        def update_status(message):
            window.after(0, status_var.set, message)
        def update_progress(percentage):
            window.after(0, progress_bar.config, {'value': percentage})
            
        result_message = editortkinter.fase_3_geracao_docx(
            config_data, 
            all_photos, 
            decisions, 
            update_status,
            update_progress
        )
        
        window.after(0, on_process_success, config_data['output_docx'])
    except Exception as e:
        window.after(0, on_process_error, e)

def on_process_success(output_doc_path):
    """GUI THREAD: Chamado ao concluir com sucesso."""
    progress_bar['value'] = 0
    progress_bar['mode'] = 'indeterminate'
    progress_bar.pack_forget() 
    run_button.config(state=tk.NORMAL)
    status_var.set("Processo concluído com sucesso!")
    # --- ADIÇÃO: Salva os caminhos atuais ---
    paths_to_save = {
        "last_input_docx": input_file_var.get(),
        "last_photos_folder": photos_folder_var.get(),
        "last_output_docx": output_file_var.get()
    }
    save_settings(paths_to_save)
    # --- FIM DA ADIÇÃO ---
    try:
        if output_doc_path and os.path.exists(output_doc_path):
            os.startfile(output_doc_path)
    except Exception as e:
        status_var.set(f"Sucesso! (Mas falhou ao abrir o doc final: {e})")

def on_process_error(error):
    """GUI THREAD: Chamado em qualquer falha."""
    progress_bar['value'] = 0
    progress_bar['mode'] = 'indeterminate'
    progress_bar.pack_forget()
    run_button.config(state=tk.NORMAL)
    status_var.set(f"ERRO: {error}")
    messagebox.showerror("Erro no Processo", f"Ocorreu um erro:\n\n{error}")


# --- GUI Principal ---
window = tk.Tk()
window.title("AutoFoto Report")
window.geometry("700x450")
window.resizable(False, False)

# Configuração de Fontes
title_font = font.Font(family="Helvetica", size=16, weight="bold")
author_font = font.Font(family="Helvetica", size=10, slant="italic")
path_font = font.Font(family="Courier", size=10)
button_font = font.Font(family="Helvetica", size=10)
main_button_font = font.Font(family="Helvetica", size=12, weight="bold")

# --- Frame Superior (Título e Botões Auxiliares) ---
top_frame = tk.Frame(window)
top_frame.pack(fill=tk.X, padx=10, pady=5)

title_subframe = tk.Frame(top_frame)
title_subframe.pack(side=tk.LEFT, expand=True)
tk.Label(title_subframe, text="AutoFoto Report", font=title_font).pack()
tk.Label(title_subframe, text="v4.2.3 | Criado por: Pedro H.G.C. Vidal", font=author_font).pack()

buttons_subframe = tk.Frame(top_frame)
buttons_subframe.pack(side=tk.RIGHT)
tk.Button(buttons_subframe, text="Configurações de Layout...", command=open_settings_window).pack(side=tk.TOP, pady=5)
tk.Button(buttons_subframe, text="Ajuda / Leia-me", command=show_help_window).pack(side=tk.TOP)

# --- Frame Principal (Seleção de Arquivos) ---
main_frame = tk.Frame(window, padx=20, pady=5)
main_frame.pack(fill=tk.BOTH, expand=True)

input_file_var = tk.StringVar(value="[Nenhum arquivo selecionado]")
photos_folder_var = tk.StringVar(value="[Nenhuma pasta selecionada]")
output_file_var = tk.StringVar(value="[Nenhum local de saída definido]")
insertion_point_var = tk.StringVar(value="[Nenhum local de inserção selecionado]")

tk.Button(main_frame, text="1. Selecionar Relatório Word...", font=button_font, command=select_input_file).pack(fill=tk.X, pady=(5,0))
tk.Label(main_frame, textvariable=input_file_var, font=path_font, relief="sunken", anchor="w", padx=5).pack(fill=tk.X)

tk.Button(main_frame, text="2. Selecionar Pasta de Fotos...", font=button_font, command=select_photos_folder).pack(fill=tk.X, pady=(15,0))
tk.Label(main_frame, textvariable=photos_folder_var, font=path_font, relief="sunken", anchor="w", padx=5).pack(fill=tk.X)

tk.Button(main_frame, text="3. Definir Local de Saída...", font=button_font, command=select_output_file).pack(fill=tk.X, pady=(15,0))
tk.Label(main_frame, textvariable=output_file_var, font=path_font, relief="sunken", anchor="w", padx=5).pack(fill=tk.X)

# --- Botão de Ponto de Inserção (Corrigido) ---
insertion_btn = tk.Button(main_frame, text="4. Escolher Local de Inserção no Word...", font=button_font, command=open_document_preview, state=tk.DISABLED)
insertion_btn.pack(fill=tk.X, pady=(15,0))
tk.Label(main_frame, textvariable=insertion_point_var, font=path_font, relief="sunken", anchor="w", padx=5, foreground="blue").pack(fill=tk.X)
# --- Fim da Correção ---

# --- Botão Principal e Barra de Progresso ---
run_button = tk.Button(window, text="Iniciar Importação e Edição de Fotos", font=main_button_font, command=run_importer, bg="#DDFFDD")
run_button.pack(pady=(20, 5)) 

progress_bar = ttk.Progressbar(window, orient="horizontal", length=300, mode="indeterminate")
# A barra será mostrada/escondida com .pack() e .pack_forget()

# --- Barra de Status ---
status_var = tk.StringVar(value="Pronto para começar.")
status_label = tk.Label(window, textvariable=status_var, relief="sunken", anchor="w", padx=10)
status_label.pack(side=tk.BOTTOM, fill=tk.X)

try:
    saved_settings = load_settings()   
    # 1. Recupera Caminho do Word
    last_input = saved_settings.get("last_input_docx", "")
    if last_input and os.path.exists(last_input):
        input_file_var.set(last_input)
        insertion_btn.config(state=tk.NORMAL) # Habilita o botão 4 se o arquivo existir    
    # 2. Recupera Pasta de Fotos
    last_photos = saved_settings.get("last_photos_folder", "")
    if last_photos and os.path.exists(last_photos):
        photos_folder_var.set(last_photos)        
    # 3. Recupera Caminho de Saída
    last_output = saved_settings.get("last_output_docx", "")
    if last_output:
        output_file_var.set(last_output)
        
except Exception as e:
    print(f"Erro ao carregar caminhos salvos: {e}")

window.mainloop()