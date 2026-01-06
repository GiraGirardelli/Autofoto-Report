import os
import sys
import json
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageTk, ImageEnhance, ImageDraw, ImageFont, ExifTags
import io
import tkinter as tk
from tkinter import messagebox, ttk

try:
    import cv2
    import numpy as np
except ImportError:
    cv2 = None

# --- CONFIGURAÇÃO ---
SCRIPT_VERSION = "3.0.2 FINAL" 
TEXTO_ANCORA = "[INSERIR_FOTOS_AQUI]"
TITLE_PREFIX = ""
TITLE_FONT_SIZE_PT = 16
TITLE_BOLD = True
CENTER_TITLES = True
CENTER_IMAGES = True
SPACE_ABOVE_TITLES_PT = 36
SPACE_ABOVE_FIRST_IMAGE_PT = 35
SPACE_BETWEEN_IMAGES_PT = 24
MAX_IMAGE_WIDTH_CM = 16.0
MAX_IMAGE_HEIGHT_CM = 22.0
FORCE_RESIZE_SMALLER_IMAGES = False

DEFAULT_SETTINGS = {
    "SCRIPT_VERSION": SCRIPT_VERSION, "TEXTO_ANCORA": TEXTO_ANCORA,
    "TITLE_PREFIX": TITLE_PREFIX, "TITLE_FONT_SIZE_PT": TITLE_FONT_SIZE_PT, 
    "TITLE_BOLD": TITLE_BOLD, "CENTER_TITLES": CENTER_TITLES, 
    "CENTER_IMAGES": CENTER_IMAGES, "SPACE_ABOVE_TITLES_PT": SPACE_ABOVE_TITLES_PT,
    "SPACE_ABOVE_FIRST_IMAGE_PT": SPACE_ABOVE_FIRST_IMAGE_PT,
    "SPACE_BETWEEN_IMAGES_PT": SPACE_BETWEEN_IMAGES_PT,
    "MAX_IMAGE_WIDTH_CM": MAX_IMAGE_WIDTH_CM,
    "MAX_IMAGE_HEIGHT_CM": MAX_IMAGE_HEIGHT_CM, 
    "FORCE_RESIZE_SMALLER_IMAGES": FORCE_RESIZE_SMALLER_IMAGES
}
# --- FIM DA CONFIGURAÇÃO ---

# --- INÍCIO DO EDITOR MESTRE ---
# --- INÍCIO DO EDITOR MESTRE ---
# (Esta classe será chamada pelo LAUNCHER, na THREAD PRINCIPAL)
# --- INÍCIO DO EDITOR MESTRE (Sem Legenda, Todos os Bugs Corrigidos) ---


##Esta é a nova classe para seleção de ordem das fotos dentro dos lotes de forma facilitada. v0.1
# --- INÍCIO DA CLASSE REORDER EDITOR (Scrollbar Corrigida) ---
class BatchReorderEditor(tk.Toplevel):
    """
    Uma janela modal para reordenar as fotos de um lote
    usando entradas numéricas (1, 2, 3...) e com pré-visualização.
    """
    def __init__(self, parent, batch_name, photo_paths_in_batch):
        super().__init__(parent)
        self.title(f"Reordenar Lote: {batch_name}")
        self.geometry("800x600")
        
        # Armazena os dados
        self.original_paths = photo_paths_in_batch
        self.result = None 
        self.entry_widgets = [] 
        self.preview_tk_image = None 

        # --- UI ---
        tk.Label(self, text="Clique no nome do arquivo para pré-visualizar.\nDigite a nova ordem (1, 2, 3...) para cada foto:", font=("Helvetica", 10, "bold")).pack(pady=10)
        
        main_pane = tk.Frame(self)
        main_pane.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # --- Painel Esquerdo (Lista de FOTOS) ---
        left_pane = tk.Frame(main_pane, width=350) 
        left_pane.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        left_pane.pack_propagate(False) # Mantém largura fixa

        list_frame_container = tk.Frame(left_pane, borderwidth=1, relief="sunken")
        list_frame_container.pack(fill=tk.BOTH, expand=True)
        
        canvas = tk.Canvas(list_frame_container)
        scrollbar = tk.Scrollbar(list_frame_container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # --- CORREÇÃO: Scrollbar empacotada PRIMEIRO ---
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        # -----------------------------------------------
        
        # --- Painel Direito (Pré-visualização) ---
        right_pane = tk.Frame(main_pane, borderwidth=1, relief="sunken")
        right_pane.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        self.preview_canvas = tk.Canvas(right_pane, bg="gray90")
        self.preview_canvas.pack(fill=tk.BOTH, expand=True)

        # Popula o frame rolável
        for i, path in enumerate(self.original_paths):
            row_frame = tk.Frame(scrollable_frame)
            row_frame.pack(fill=tk.X, padx=5, pady=3)
            
            entry_var = tk.StringVar(value=str(i + 1))
            entry = tk.Entry(row_frame, textvariable=entry_var, width=4, font=("Helvetica", 11))
            entry.pack(side=tk.LEFT, padx=5)
            self.entry_widgets.append(entry)
            
            file_name = os.path.basename(path)
            filename_label = tk.Label(row_frame, text=file_name, anchor="w", font=("Courier", 10, "underline"), fg="blue", cursor="hand2")
            filename_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            filename_label.bind("<Button-1>", lambda event, idx=i: self._show_preview(idx))
            
        # Botões de Ação
        button_frame = tk.Frame(self)
        button_frame.pack(pady=10)
        
        tk.Button(button_frame, text="Cancelar", command=self.destroy).pack(side=tk.RIGHT, padx=10)
        tk.Button(button_frame, text="Confirmar Ordem", command=self._on_confirm, font=("Helvetica", 10, "bold")).pack(side=tk.RIGHT)

        self.transient(parent)
        self.grab_set()
        self.wait_window()

    def _show_preview(self, index):
        """Carrega e exibe a imagem clicada no canvas de pré-visualização."""
        path = self.original_paths[index]
        try:
            pil_image = Image.open(path)
            
            if hasattr(pil_image, '_getexif'):
                exif = pil_image._getexif()
                if exif:
                    orientation = exif.get(0x0112, 1) # 1 = Normal
                    if orientation == 3:
                        pil_image = pil_image.transpose(Image.ROTATE_180)
                    elif orientation == 6:
                        pil_image = pil_image.transpose(Image.ROTATE_270)
                    elif orientation == 8:
                        pil_image = pil_image.transpose(Image.ROTATE_90)
            
            self.preview_canvas.update_idletasks() 
            canvas_w = self.preview_canvas.winfo_width()
            canvas_h = self.preview_canvas.winfo_height()
            
            if canvas_w < 50 or canvas_h < 50:
                canvas_w, canvas_h = 400, 400
                
            pil_image.thumbnail((canvas_w - 2, canvas_h - 2))
            
            self.preview_tk_image = ImageTk.PhotoImage(pil_image)
            
            self.preview_canvas.delete("all")
            self.preview_canvas.create_image(
                canvas_w / 2, canvas_h / 2, 
                anchor=tk.CENTER, 
                image=self.preview_tk_image
            )
            
        except Exception as e:
            self.preview_canvas.delete("all")
            canvas_w = self.preview_canvas.winfo_width()
            self.preview_canvas.create_text(
                10, 10, anchor=tk.NW, 
                text=f"Erro ao carregar pré-visualização:\n{e}", 
                fill="red", width=canvas_w - 20
            )

    def _on_confirm(self):
        """Valida a entrada e prepara a lista de resultados."""
        num_items = len(self.original_paths)
        new_order_map = []
        
        try:
            new_indices = [int(entry.get()) for entry in self.entry_widgets]
        except ValueError:
            messagebox.showerror("Erro de Validação", "Todas as ordens devem ser números inteiros (ex: 1, 2, 3).", parent=self)
            return

        if len(set(new_indices)) != num_items:
            messagebox.showerror("Erro de Validação", "Existem números de ordem duplicados. Cada foto deve ter um número único.", parent=self)
            return
            
        expected_set = set(range(1, num_items + 1))
        if set(new_indices) != expected_set:
            messagebox.showerror("Erro de Validação", f"Os números de ordem devem ser uma sequência de 1 a {num_items}, sem lacunas ou números maiores.", parent=self)
            return

        for i, new_idx in enumerate(new_indices):
            original_path = self.original_paths[i]
            new_order_map.append((new_idx - 1, original_path)) 
            
        new_order_map.sort(key=lambda item: item[0])
        self.result = [path for (idx, path) in new_order_map]
        self.destroy()

    def get_new_order(self):
        """Método público para o MasterCropEditor pegar o resultado."""
        return self.result
# --- FIM DA CLASSE REORDER EDITOR ---


# --- INÍCIO DO EDITOR MESTRE ---
class MasterCropEditor(tk.Toplevel):
    def __init__(self, parent, flat_photo_list, all_photos_by_lote, callback_on_finish):
        super().__init__(parent)
        self.state('zoomed') 
        self.title("AutoFoto Report: Editor de fotos")
        
        self.photo_list = flat_photo_list
        self.callback_on_finish = callback_on_finish
        
        self.decisions = {} 
        self.current_index = -1
        self.last_relative_crop = None

        self.batch_boundaries = {} 
        current_idx = 0
        sorted_batch_names = sorted(all_photos_by_lote.keys()) 
        for batch_name in sorted_batch_names:
            photos_in_batch = all_photos_by_lote[batch_name]
            start = current_idx
            end = current_idx + len(photos_in_batch)
            self.batch_boundaries[batch_name] = (start, end) 
            current_idx = end

        self.brightness_timer = None
        self.original_pil_image = None
        self.crop_rect = None
        self.crop_coords = {}
        self.pil_thumb = None
        self.last_crop_overlay = None
        self.tk_image_cache = {}

        # --- Layout da GUI ---
        main_frame = tk.Frame(self)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # --- Frame Esquerdo ---
        left_frame = tk.Frame(main_frame, width=300)
        left_frame.pack(side=tk.LEFT, fill=tk.Y)
        
        tk.Label(left_frame, text="Ordem das Fotos:", font=("Helvetica", 12, "bold")).pack(anchor="w")
        
        move_buttons_frame = tk.Frame(left_frame)
        move_buttons_frame.pack(fill=tk.X, pady=5)
        self.move_up_btn = tk.Button(move_buttons_frame, text="Mover Cima ↑", command=self._on_move_up, state=tk.DISABLED)
        self.move_up_btn.pack(side=tk.LEFT, expand=True, fill=tk.X)
        self.move_down_btn = tk.Button(move_buttons_frame, text="Mover Baixo ↓", command=self._on_move_down, state=tk.DISABLED)
        self.move_down_btn.pack(side=tk.RIGHT, expand=True, fill=tk.X)
        
        self.reorder_batch_btn = tk.Button(left_frame, text="Reordenar Lote (1,2,3...)", command=self._open_reorder_editor, state=tk.DISABLED)
        self.reorder_batch_btn.pack(fill=tk.X, pady=(5,0))
        
        list_frame = tk.Frame(left_frame)
        list_frame.pack(fill=tk.BOTH, expand=True)
        list_scrollbar = tk.Scrollbar(list_frame, orient=tk.VERTICAL)
        self.listbox = tk.Listbox(list_frame, yscrollcommand=list_scrollbar.set, exportselection=False, font=("Courier", 9))
        list_scrollbar.config(command=self.listbox.yview)
        list_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Atalhos Visuais
        shortcut_frame = tk.LabelFrame(left_frame, text="Atalhos do Teclado", font=("Helvetica", 9, "bold"))
        shortcut_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(10, 0), ipady=2)
        shortcut_text = (
            "Aplicar/Próximo:  Enter\n"
            "Pular/Próximo:    Seta Direita\n"
            "Girar Esq/Dir:    Ctrl + Q / Ctrl + W\n"
            "Mover Foto:       Ctrl + Seta Cima/Baixo\n"
            "Aplicar Último:   Ctrl + L\n"
            "Replicar Demais:  Ctrl + R"
        )
        tk.Label(shortcut_frame, text=shortcut_text, font=("Courier", 9), justify=tk.LEFT).pack(anchor="w", padx=5, pady=2)

        for i, path in enumerate(self.photo_list):
            dir_name, file_name = os.path.split(path)
            lote_name = os.path.basename(dir_name)
            self.listbox.insert(tk.END, f"{lote_name} / {file_name}")
            self.decisions[path] = {"crop": "skip", "rotation": 0, "brightness": 1.0}

        self.listbox.bind("<<ListboxSelect>>", self._on_list_select)

        # --- Frame Direito ---
        right_frame = tk.Frame(main_frame)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        self.canvas = tk.Canvas(right_frame, bg="gray90", highlightthickness=0)
        self.canvas.pack(fill=tk.BOTH, expand=True)
        
        self.canvas.bind("<ButtonPress-1>", self._on_canvas_press)
        self.canvas.bind("<B1-Motion>", self._on_canvas_drag)

        brightness_frame = tk.Frame(right_frame)
        brightness_frame.pack(fill=tk.X, pady=5)
        tk.Label(brightness_frame, text="Luminosidade:").pack(side=tk.LEFT, padx=5)
        self.brightness_slider = tk.Scale(
            brightness_frame, from_=0.5, to=1.5, resolution=0.05, 
            orient=tk.HORIZONTAL, command=self._on_brightness_change
        )
        self.brightness_slider.set(1.0)
        self.brightness_slider.pack(fill=tk.X, expand=True, side=tk.LEFT)
        self.reset_brightness_btn = tk.Button(brightness_frame, text="Reset", command=self._reset_brightness)
        self.reset_brightness_btn.pack(side=tk.RIGHT, padx=5)
        
        # Botões de Rotação em Lote
        batch_rot_frame = tk.Frame(right_frame)
        batch_rot_frame.pack(fill=tk.X, pady=(10, 0))
        tk.Label(batch_rot_frame, text="Girar Todas:").pack(side=tk.LEFT, padx=5)
        tk.Button(batch_rot_frame, text="Tudo Horizontal ⬌", command=lambda: self._batch_rotate('landscape')).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=2)
        tk.Button(batch_rot_frame, text="Tudo Vertical ⬍", command=lambda: self._batch_rotate('portrait')).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=2)

        self.dimensions_var = tk.StringVar(value="Selecione uma foto para começar.")
        tk.Label(right_frame, textvariable=self.dimensions_var, anchor="w").pack(fill=tk.X)

        # --- Frame Inferior ---
        bottom_frame = tk.Frame(self, height=60, pady=10, padx=10)
        bottom_frame.pack(fill=tk.X, side=tk.BOTTOM)
        
        self.finish_btn = tk.Button(bottom_frame, text="1. Finalizar Edição", command=self._on_finish, bg="#E0FFE0", font=("Helvetica", 10, "bold"))
        self.finish_btn.pack(side=tk.LEFT, padx=5)
        self.rotate_left_btn = tk.Button(bottom_frame, text="Girar 90° Esq", command=self._on_rotate_left)
        self.rotate_left_btn.pack(side=tk.LEFT, padx=5)
        self.rotate_right_btn = tk.Button(bottom_frame, text="Girar 90° Dir", command=self._on_rotate_right)
        self.rotate_right_btn.pack(side=tk.LEFT, padx=2)
        self.replicate_btn = tk.Button(bottom_frame, text="Replicar para Demais", command=self._on_replicate, state=tk.DISABLED)
        self.replicate_btn.pack(side=tk.RIGHT, padx=5)
        self.apply_last_btn = tk.Button(bottom_frame, text="Aplicar Último Corte", command=self._on_apply_last, state=tk.DISABLED)
        self.apply_last_btn.pack(side=tk.RIGHT, padx=5)
        self.skip_btn = tk.Button(bottom_frame, text="Pular (Usar Original) e Próximo >>", command=self._on_skip)
        self.skip_btn.pack(side=tk.RIGHT, padx=5)
        self.apply_btn = tk.Button(bottom_frame, text="Aplicar Corte e Próximo >>", command=self._on_apply)
        self.apply_btn.pack(side=tk.RIGHT, padx=5)

        # --- Inicialização ---
        self.lift()
        self.focus_force()
        self.grab_set()
        self.canvas.update_idletasks() 
        self._select_next()
        
        # Configura Atalhos
        self._setup_key_binds()
        
        self.protocol("WM_DELETE_WINDOW", self._on_finish)

    # --- Setup de Atalhos ---
    def _is_typing(self, event):
        return isinstance(event.widget, tk.Entry)
    def _on_key_apply(self, event):
        if self._is_typing(event): return
        self._on_apply()
    def _on_key_skip(self, event):
        if self._is_typing(event): return
        self._on_skip()
    def _on_key_rotate_left(self, event):
        if self._is_typing(event): return
        self._on_rotate_left()
    def _on_key_rotate_right(self, event):
        if self._is_typing(event): return
        self._on_rotate_right()
    def _on_key_move_up(self, event):
        if self._is_typing(event): return
        self._on_move_up()
    def _on_key_move_down(self, event):
        if self._is_typing(event): return
        self._on_move_down()
    def _on_key_apply_last(self, event):
        if self._is_typing(event): return
        self._on_apply_last()
    def _on_key_replicate(self, event):
        if self._is_typing(event): return
        self._on_replicate()

    def _setup_key_binds(self):
        self.listbox.bind("<Return>", self._on_key_apply)
        self.listbox.bind("<Right>", self._on_key_skip)
        self.listbox.bind("<Control-q>", self._on_key_rotate_left)
        self.listbox.bind("<Control-w>", self._on_key_rotate_right)
        self.listbox.bind("<Control-Up>", self._on_key_move_up)
        self.listbox.bind("<Control-Down>", self._on_key_move_down)
        self.listbox.bind("<Control-l>", self._on_key_apply_last)
        self.listbox.bind("<Control-r>", self._on_key_replicate)

    # --- Lógica de Rotação em Lote ---
    def _batch_rotate(self, target_mode):
        count = 0
        for path in self.photo_list:
            try:
                img = Image.open(path)
                w, h = img.size
                # Ignora EXIF para consistência visual
                
                current_manual_rotation = self.decisions[path].get("rotation", 0)
                if current_manual_rotation in [90, 270]:
                    w, h = h, w
                    
                is_currently_landscape = w > h
                needs_rotation = False
                
                if target_mode == 'landscape' and not is_currently_landscape:
                    needs_rotation = True
                elif target_mode == 'portrait' and is_currently_landscape:
                    needs_rotation = True
                    
                if needs_rotation:
                    new_rotation = (current_manual_rotation + 90) % 360
                    self.decisions[path]["rotation"] = new_rotation
                    count += 1
            except Exception as e:
                print(f"Erro batch rotate: {e}")

        if count > 0:
            self._update_canvas_image()
            # --- CORREÇÃO AQUI: parent=self ---
            messagebox.showinfo("Rotação em Lote", f"{count} fotos foram giradas.", parent=self)
        else:
            # --- CORREÇÃO AQUI: parent=self ---
            messagebox.showinfo("Rotação em Lote", "Nenhuma rotação necessária.", parent=self)

    # --- Controle ---
    def _select_next(self):
        new_index = self.current_index + 1
        if new_index < len(self.photo_list):
            self.listbox.selection_clear(0, tk.END)
            self.listbox.selection_set(new_index)
            self.listbox.activate(new_index)
            self.listbox.see(new_index)
            self._load_image_at_index(new_index)
            self._update_list_buttons()
        else:
            self.dimensions_var.set("Todas as fotos foram processadas! Finalizando...")
            self.after(1000, self._on_finish)

    def _on_brightness_change(self, value_str):
        if self.brightness_timer:
            self.after_cancel(self.brightness_timer)
        self.brightness_timer = self.after(
            150, lambda v=value_str: self._apply_brightness_filter(v)
        )

    def _apply_brightness_filter(self, value_str):
        if self.current_index < 0: return 
        value = float(value_str)
        path = self.photo_list[self.current_index] 
        self.decisions[path]["brightness"] = value
        self._update_canvas_image()

    def _reset_brightness(self):
        if self.brightness_timer:
            self.after_cancel(self.brightness_timer)
            self.brightness_timer = None
        self.brightness_slider.set(1.0)
        self._apply_brightness_filter("1.0")

    def _update_canvas_image(self):
        if not self.original_pil_image: return
        path = self.photo_list[self.current_index]
        dec = self.decisions.get(path, {})
        rotation_state = dec.get("rotation", 0)
        brightness_state = dec.get("brightness", 1.0)

        temp_thumb = self.original_pil_image.copy()

        if rotation_state == 90:
            temp_thumb = temp_thumb.transpose(Image.ROTATE_270)
        elif rotation_state == 180:
            temp_thumb = temp_thumb.transpose(Image.ROTATE_180)
        elif rotation_state == 270:
            temp_thumb = temp_thumb.transpose(Image.ROTATE_90)

        if brightness_state != 1.0:
            enhancer = ImageEnhance.Brightness(temp_thumb)
            temp_thumb = enhancer.enhance(brightness_state)

        self.pil_thumb = temp_thumb
        canvas_w, canvas_h = self.canvas.winfo_width(), self.canvas.winfo_height()
        if canvas_w < 50 or canvas_h < 50: canvas_w, canvas_h = 800, 600
        self.pil_thumb.thumbnail((canvas_w - 2, canvas_h - 2))

        new_tk_image = ImageTk.PhotoImage(self.pil_thumb)
        self.tk_image_cache[path] = new_tk_image

        self.canvas.delete("all")
        self.canvas.create_image(0, 0, anchor=tk.NW, image=new_tk_image)
        self.canvas.image = new_tk_image

        self.dimensions_var.set(f"Arquivo: {os.path.basename(path)} ({self.pil_thumb.width} x {self.pil_thumb.height} px)")
        
        self._draw_last_crop_overlay()
        if self.crop_rect:
            self.crop_rect = self.canvas.create_rectangle(
                self.crop_coords.get('start_x', 0), self.crop_coords.get('start_y', 0),
                self.crop_coords.get('end_x', 0), self.crop_coords.get('end_y', 0),
                outline='green', width=2
            )
        else:
            self.crop_rect = None

    def _on_rotate_left(self):
        self._apply_rotation(-90)
    def _on_rotate_right(self):
        self._apply_rotation(90)
    def _apply_rotation(self, angle):
        if self.current_index < 0: return
        path = self.photo_list[self.current_index]
        current_rotation = self.decisions[path].get("rotation", 0)
        new_rotation = (current_rotation + angle) % 360
        self.decisions[path]["rotation"] = new_rotation
        self._update_canvas_image()
        
    def _on_list_select(self, event):
        try:
            selected_indices = self.listbox.curselection()
            if not selected_indices: return
            self._load_image_at_index(selected_indices[0])
            self._update_list_buttons()
        except Exception as e:
            # --- CORREÇÃO AQUI: parent=self ---
            messagebox.showerror("Erro", f"Erro ao carregar: {e}", parent=self)

    def _load_image_at_index(self, index):
        self.current_index = index
        path = self.photo_list[index]
        try:
            dec = self.decisions.get(path, {"rotation": 0, "brightness": 1.0})
            rotation_state = dec.get("rotation", 0)
            brightness_state = dec.get("brightness", 1.0)

            self.original_pil_image = Image.open(path)
            self.brightness_slider.set(brightness_state)
            self.crop_rect = None
            self.crop_coords = {}

            self._update_canvas_image()

            if self.last_relative_crop:
                self.apply_last_btn.config(state=tk.NORMAL)
                self.replicate_btn.config(state=tk.NORMAL)
            else:
                self.apply_last_btn.config(state=tk.DISABLED)
                self.replicate_btn.config(state=tk.DISABLED)
        except Exception as e:
            self.canvas.delete("all")
            self.canvas.create_text(10, 10, anchor="nw", text=f"Erro: {e}", fill="red")
            
    def _draw_last_crop_overlay(self):
        if self.last_crop_overlay: self.canvas.delete(self.last_crop_overlay)
        if self.last_relative_crop:
            w, h = self.pil_thumb.width, self.pil_thumb.height
            c = self.last_relative_crop
            x1, y1 = round(c['x1'] * w), round(c['y1'] * h)
            x2, y2 = round(c['x2'] * w), round(c['y2'] * h)
            self.last_crop_overlay = self.canvas.create_rectangle(x1, y1, x2, y2, outline='blue', width=2, dash=(5, 2))

    def _on_canvas_press(self, event):
        if self.crop_rect: self.canvas.delete(self.crop_rect)
        w, h = self.pil_thumb.width, self.pil_thumb.height 
        x, y = max(0, min(event.x, w)), max(0, min(event.y, h))
        self.crop_coords = {'start_x': x, 'start_y': y}
        self.crop_rect = self.canvas.create_rectangle(x, y, x, y, outline='green', width=2)
        self.listbox.focus_set() # Devolve foco para atalhos
        
    def _on_canvas_drag(self, event):
        if not self.crop_rect or 'start_x' not in self.crop_coords: return
        w, h = self.pil_thumb.width, self.pil_thumb.height
        x, y = max(0, min(event.x, w)), max(0, min(event.y, h))
        self.crop_coords['end_x'], self.crop_coords['end_y'] = x, y
        self.canvas.coords(self.crop_rect, self.crop_coords['start_x'], self.crop_coords['start_y'], x, y)
        w_crop = abs(x - self.crop_coords['start_x'])
        h_crop = abs(y - self.crop_coords['start_y'])
        self.dimensions_var.set(f"Seleção: {w_crop} x {h_crop} px")

    def _save_current_selection(self):
        if not self.crop_coords or 'end_x' not in self.crop_coords:
            # --- CORREÇÃO AQUI: parent=self ---
            messagebox.showwarning("Sem Seleção", "Desenhe um retângulo primeiro.", parent=self)
            return None
        w, h = self.pil_thumb.width, self.pil_thumb.height 
        x1 = min(self.crop_coords['start_x'], self.crop_coords['end_x'])
        y1 = min(self.crop_coords['start_y'], self.crop_coords['end_y'])
        x2 = max(self.crop_coords['start_x'], self.crop_coords['end_x'])
        y2 = max(self.crop_coords['start_y'], self.crop_coords['end_y'])
        if x1 >= x2 or y1 >= y2:
            # --- CORREÇÃO AQUI: parent=self ---
            messagebox.showwarning("Inválido", "Seleção com tamanho 0.", parent=self)
            return None
        relative_crop = {'x1': x1/w, 'y1': y1/h, 'x2': x2/w, 'y2': y2/h}
        self.last_relative_crop = relative_crop
        return relative_crop

    def _on_apply(self):
        relative_crop = self._save_current_selection()
        if relative_crop:
            path = self.photo_list[self.current_index]
            self.decisions[path]["crop"] = relative_crop
            self.listbox.itemconfig(self.current_index, {'bg': "#DFFFE0"})
            self._select_next()

    def _on_skip(self):
        path = self.photo_list[self.current_index]
        self.decisions[path]["crop"] = "skip"
        self.listbox.itemconfig(self.current_index, {'bg': '#FFFADF'})
        self._select_next()

    def _on_apply_last(self):
        if not self.last_relative_crop: return
        path = self.photo_list[self.current_index]
        self.decisions[path]["crop"] = self.last_relative_crop
        self.listbox.itemconfig(self.current_index, {'bg': '#DFFFE0'})
        self._select_next()
        
    def _on_replicate(self):
        relative_crop = self._save_current_selection()
        current_rotation = 0
        current_brightness = 1.0
        if self.current_index >= 0:
            path = self.photo_list[self.current_index]
            current_rotation = self.decisions[path].get("rotation", 0)
            current_brightness = self.decisions[path].get("brightness", 1.0)

        if not relative_crop:
            if self.last_relative_crop:
                # --- CORREÇÃO AQUI: parent=self ---
                if messagebox.askyesno("Confirmar", "Replicar último corte, brilho e rotação?", parent=self):
                    relative_crop = self.last_relative_crop
                else:
                    return
            else:
                return 
                
        for index in range(self.current_index, len(self.photo_list)):
            path = self.photo_list[index]
            self.decisions[path]["crop"] = relative_crop
            self.decisions[path]["rotation"] = current_rotation
            self.decisions[path]["brightness"] = current_brightness
            self.listbox.itemconfig(index, {'bg': '#DFFFE0'})
            
        self.dimensions_var.set(f"Configurações replicadas para o restante.")
        self.after(1000, self._on_finish)

    # --- Métodos de Reordenação ---
    def _open_reorder_editor(self):
        selected_indices = self.listbox.curselection()
        if not selected_indices: return
        idx = selected_indices[0]
        path = self.photo_list[idx]
        batch_name = os.path.basename(os.path.dirname(path)) 
        start_idx, end_idx = self._get_batch_limits(idx)
        
        if start_idx is None:
            # --- CORREÇÃO AQUI: parent=self ---
            messagebox.showerror("Erro", "Limites de lote não encontrados.", parent=self)
            return

        paths_to_reorder = self.photo_list[start_idx:end_idx]
        editor = BatchReorderEditor(self, batch_name, paths_to_reorder)
        newly_ordered_paths = editor.get_new_order()
        
        if newly_ordered_paths:
            current_path_selected = self.photo_list[idx]
            self.photo_list[start_idx:end_idx] = newly_ordered_paths
            self._rebuild_listbox()
            try:
                new_selection_idx = self.photo_list.index(current_path_selected)
                self.listbox.selection_set(new_selection_idx)
                self.listbox.activate(new_selection_idx)
                self.listbox.see(new_selection_idx)
                self.current_index = new_selection_idx
            except ValueError:
                self.listbox.selection_set(start_idx) 
                self.current_index = start_idx
            self._update_list_buttons()

    def _rebuild_listbox(self):
        self.listbox.delete(0, tk.END)
        for i, path in enumerate(self.photo_list):
            dir_name, file_name = os.path.split(path)
            lote_name = os.path.basename(dir_name)
            self.listbox.insert(tk.END, f"{lote_name} / {file_name}")
            
            decision = self.decisions.get(path, {"crop": "skip"})
            crop_status = decision.get("crop", "skip")
            
            if crop_status == "skip":
                self.listbox.itemconfig(i, {'bg': '#FFFADF'})
            elif isinstance(crop_status, dict):
                self.listbox.itemconfig(i, {'bg': '#DFFFE0'})
    
    def _get_batch_limits(self, index):
        if index < 0 or index >= len(self.photo_list):
            return None, None
        path = self.photo_list[index]
        batch_name = os.path.basename(os.path.dirname(path)) 
        
        current_batch_indices = [
            i for i, p in enumerate(self.photo_list) 
            if os.path.basename(os.path.dirname(p)) == batch_name
        ]
        if not current_batch_indices:
            return None, None
        
        current_start = min(current_batch_indices)
        current_end = max(current_batch_indices) + 1
        
        if current_start <= index < current_end:
             return current_start, current_end
        else:
             return None, None

    def _update_list_buttons(self):
        selected_indices = self.listbox.curselection()
        if not selected_indices:
            self.move_up_btn.config(state=tk.DISABLED)
            self.move_down_btn.config(state=tk.DISABLED)
            self.reorder_batch_btn.config(state=tk.DISABLED)
            return

        idx = selected_indices[0]
        start_idx, end_idx = self._get_batch_limits(idx)

        self.reorder_batch_btn.config(state=tk.NORMAL)
        
        if start_idx is None: 
             self.move_up_btn.config(state=tk.DISABLED)
             self.move_down_btn.config(state=tk.DISABLED)
             return

        if idx > start_idx:
            self.move_up_btn.config(state=tk.NORMAL)
        else:
            self.move_up_btn.config(state=tk.DISABLED)

        if idx < end_idx - 1:
            self.move_down_btn.config(state=tk.NORMAL)
        else:
            self.move_down_btn.config(state=tk.DISABLED)
            
    def _on_move_up(self):
        self._move_item(-1)
    def _on_move_down(self):
        self._move_item(1)
    def _move_item(self, direction):
        selected_indices = self.listbox.curselection()
        if not selected_indices: return
        idx = selected_indices[0]
        new_idx = idx + direction
        start_idx, end_idx = self._get_batch_limits(idx)
        if start_idx is None or not (start_idx <= new_idx < end_idx):
            return

        self.photo_list[idx], self.photo_list[new_idx] = self.photo_list[new_idx], self.photo_list[idx]
        
        bg_colors = [self.listbox.itemcget(i, "bg") for i in range(len(self.photo_list))]
        bg_colors[idx], bg_colors[new_idx] = bg_colors[new_idx], bg_colors[idx]
        
        self.listbox.delete(0, tk.END)
        for i, path in enumerate(self.photo_list):
            dir_name, file_name = os.path.split(path)
            lote_name = os.path.basename(dir_name)
            self.listbox.insert(tk.END, f"{lote_name} / {file_name}")
            if bg_colors[i]:
                self.listbox.itemconfig(i, {'bg': bg_colors[i]})

        self.listbox.selection_set(new_idx)
        self.listbox.activate(new_idx)
        self.listbox.see(new_idx)
        self._update_list_buttons()
        self.current_index = new_idx
        
    def _on_finish(self):
        reordered_list = self.photo_list 
        decisions_to_return = self.decisions
        
        self.grab_release()
        self.destroy()
        self.callback_on_finish(reordered_list, decisions_to_return) 

# --- FIM DO EDITOR MESTRE ---


# --- FUNÇÕES AUXILIARES ---
def find_anchor_paragraph(doc, anchor_text):
    for p in doc.paragraphs:
        if anchor_text in p.text: return p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if anchor_text in p.text: return p
    return None

def remove_paragraph(paragraph):
    p_element = paragraph._element; p_element.getparent().remove(p_element)

def get_sorted_subfolders(path):
    return sorted([f for f in os.listdir(path) if os.path.isdir(os.path.join(path, f))])

def get_sorted_images(path):
    images = [img for img in os.listdir(path) if img.lower().endswith(('.png', '.jpg', '.jpeg'))]
    try: images.sort(key=lambda x: int(os.path.splitext(x)[0]))
    except ValueError: images.sort()
    return images
# --- FIM DAS FUNÇÕES AUXILIARES ---


# --- INÍCIO DAS FUNÇÕES DE PROCESSAMENTO (REFATORADAS) ---
def fase_1_analise(pasta_principal_fotos, status_callback):
    """
    Fase 1: Analisa pastas (Modo único ou lotes).
    """
    status_callback("Analisando pastas de fotos...")
    all_photos_by_lote = {} 
    flat_photo_list = []
    total_photos = 0
    
    # 1. Tenta encontrar imagens diretamente na pasta principal
    imagens_na_raiz = get_sorted_images(pasta_principal_fotos)
    
    if imagens_na_raiz:
        # MODO PASTA ÚNICA ATIVADO
        status_callback("Modo de Pasta Única detectado.")
        lote_nome = os.path.basename(pasta_principal_fotos)
        
        paths = [os.path.join(pasta_principal_fotos, img) for img in imagens_na_raiz]
        all_photos_by_lote[lote_nome] = paths
        flat_photo_list.extend(paths)
        total_photos = len(imagens_na_raiz)
        
        status_callback(f"Encontradas {total_photos} fotos em 1 lote.")
        return all_photos_by_lote, flat_photo_list
    
    # 2. Se não achou imagens na raiz, procura por subpastas (MODO LOTES)
    status_callback("Modo de Lotes (subpastas) detectado.")
    lotes = get_sorted_subfolders(pasta_principal_fotos)
    
    if not lotes:
        raise Exception(f"Nenhuma imagem encontrada em '{pasta_principal_fotos}' E nenhuma subpasta (lote) foi encontrada.")

    for lote in lotes:
        caminho_lote = os.path.join(pasta_principal_fotos, lote)
        imagens_no_lote = get_sorted_images(caminho_lote) 
        
        if imagens_no_lote:
            paths = [os.path.join(caminho_lote, img) for img in imagens_no_lote]
            all_photos_by_lote[lote] = paths
            flat_photo_list.extend(paths)
            total_photos += len(imagens_no_lote)
            
    if total_photos == 0:
        raise Exception("Nenhuma foto encontrada nas subpastas (lotes).")
        
    status_callback(f"Encontradas {total_photos} fotos em {len(lotes)} lotes.")
    return all_photos_by_lote, flat_photo_list

def fase_3_geracao_docx(config_data, all_photos_by_lote, final_crop_decisions, status_callback, progress_callback):
    """
    Fase 3: Gera o documento.
    *** VERSÃO COM CARIMBO DE DATA/HORA ***
    """
    
    status_callback("Iniciando geração do documento...")
    try:
        caminho_documento_original = config_data['input_docx']
        caminho_documento_final = config_data['output_docx']
        layout = config_data.get('layout', DEFAULT_SETTINGS)
        paragraph_index = config_data['paragraph_index']
    except KeyError:
        raise Exception("Erro ao ler a configuração: 'paragraph_index' não foi definido.")
    except Exception as e:
        raise Exception(f"Erro ao ler a configuração: {e}")

    # Configurações
    SCRIPT_VERSION = layout.get("SCRIPT_VERSION", DEFAULT_SETTINGS["SCRIPT_VERSION"])
    TITLE_PREFIX = layout.get("TITLE_PREFIX", DEFAULT_SETTINGS["TITLE_PREFIX"])
    TITLE_FONT_SIZE_PT = layout.get("TITLE_FONT_SIZE_PT", DEFAULT_SETTINGS["TITLE_FONT_SIZE_PT"])
    TITLE_BOLD = layout.get("TITLE_BOLD", DEFAULT_SETTINGS["TITLE_BOLD"])
    CENTER_TITLES = layout.get("CENTER_TITLES", DEFAULT_SETTINGS["CENTER_TITLES"])
    CENTER_IMAGES = layout.get("CENTER_IMAGES", DEFAULT_SETTINGS["CENTER_IMAGES"])
    SPACE_ABOVE_TITLES_PT = layout.get("SPACE_ABOVE_TITLES_PT", DEFAULT_SETTINGS["SPACE_ABOVE_TITLES_PT"])
    SPACE_ABOVE_FIRST_IMAGE_PT = layout.get("SPACE_ABOVE_FIRST_IMAGE_PT", DEFAULT_SETTINGS["SPACE_ABOVE_FIRST_IMAGE_PT"])
    SPACE_BETWEEN_IMAGES_PT = layout.get("SPACE_BETWEEN_IMAGES_PT", DEFAULT_SETTINGS["SPACE_BETWEEN_IMAGES_PT"])
    MAX_IMAGE_WIDTH_CM = layout.get("MAX_IMAGE_WIDTH_CM", DEFAULT_SETTINGS["MAX_IMAGE_WIDTH_CM"])
    MAX_IMAGE_HEIGHT_CM = layout.get("MAX_IMAGE_HEIGHT_CM", DEFAULT_SETTINGS["MAX_IMAGE_HEIGHT_CM"])
    FORCE_RESIZE_SMALLER_IMAGES = layout.get("FORCE_RESIZE_SMALLER_IMAGES", DEFAULT_SETTINGS["FORCE_RESIZE_SMALLER_IMAGES"])
    SIDE_BY_SIDE_LAYOUT = layout.get("SIDE_BY_SIDE_LAYOUT", False)
    
    # --- NOVA CONFIGURAÇÃO ---
    ADD_TIMESTAMP = layout.get("ADD_TIMESTAMP", False)

    if cv2 is None: raise Exception("OpenCV (cv2) não encontrado!")
    if not caminho_documento_original or not os.path.exists(caminho_documento_original):
        raise Exception(f"Arquivo DOCX não encontrado em: {caminho_documento_original}")

    status_callback("Carregando documento Word...")
    doc = Document(caminho_documento_original)
    
    try:
        if paragraph_index >= len(doc.paragraphs):
             raise Exception(f"Índice de parágrafo fora dos limites.")
        cursor = doc.paragraphs[paragraph_index]._element
    except Exception as e:
        raise Exception(f"Erro ao encontrar o ponto de inserção: {e}")

    # Limites
    page_section = doc.sections[0]
    page_width_cm = page_section.page_width.cm
    left_margin_cm = page_section.left_margin.cm
    right_margin_cm = page_section.right_margin.cm
    page_safety_width_cm = page_width_cm - left_margin_cm - right_margin_cm
    
    page_height_cm = page_section.page_height.cm
    top_margin_cm = page_section.top_margin.cm
    bottom_margin_cm = page_section.bottom_margin.cm
    page_safety_height_cm = page_height_cm - top_margin_cm - bottom_margin_cm

    if SIDE_BY_SIDE_LAYOUT:
        limit_width_cm = min(MAX_IMAGE_WIDTH_CM, (page_safety_width_cm / 2) - 0.5)
    else:
        limit_width_cm = min(MAX_IMAGE_WIDTH_CM, page_safety_width_cm)
    
    limit_height_cm = min(MAX_IMAGE_HEIGHT_CM, page_safety_height_cm)
    
    # Total fotos
    total_fotos_a_processar = sum(len(paths) for paths in all_photos_by_lote.values())
    fotos_processadas = 0

    total_lotes = len(all_photos_by_lote)
    for idx_lote, (lote, image_paths) in enumerate(all_photos_by_lote.items()):
        status_callback(f"Processando Lote {idx_lote + 1}/{total_lotes}: {lote}")
        
        # Título do Lote
        texto_do_titulo = f"{TITLE_PREFIX}{lote}"
        titulo_paragrafo = doc.add_paragraph()
        titulo_run = titulo_paragrafo.add_run(texto_do_titulo)
        titulo_font, titulo_format = titulo_run.font, titulo_paragrafo.paragraph_format
        titulo_font.size, titulo_font.bold = Pt(TITLE_FONT_SIZE_PT), TITLE_BOLD
        titulo_format.space_before = Pt(SPACE_ABOVE_TITLES_PT)
        if CENTER_TITLES: titulo_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_format.keep_with_next = True
        
        cursor.addnext(titulo_paragrafo._element) 
        cursor = titulo_paragrafo._element
        
        # Tabela (se necessário)
        current_table = None
        if SIDE_BY_SIDE_LAYOUT and len(image_paths) > 0:
            current_table = doc.add_table(rows=0, cols=2)
            current_table.autofit = True
            cursor.addnext(current_table._element)
            cursor = current_table._element
        
        total_imagens_no_lote = len(image_paths)
        for i, caminho_imagem_completo in enumerate(image_paths):
            
            target_run = None
            target_paragraph = None

            if SIDE_BY_SIDE_LAYOUT:
                row_idx = i // 2 
                col_idx = i % 2  
                if col_idx == 0:
                    current_row = current_table.add_row()
                    current_row.cells[0].paragraphs[0].paragraph_format.keep_with_next = True
                else:
                    current_row = current_table.rows[-1]
                
                target_cell = current_row.cells[col_idx]
                target_paragraph = target_cell.paragraphs[0]
                if CENTER_IMAGES: target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                target_run = target_paragraph.add_run()
            else:
                if i > 0 and SPACE_BETWEEN_IMAGES_PT > 0:
                    spacing_paragraph = doc.add_paragraph()
                    spacing_paragraph.add_run().font.size = Pt(SPACE_BETWEEN_IMAGES_PT)
                    cursor.addnext(spacing_paragraph._element)
                    cursor = spacing_paragraph._element

                target_paragraph = doc.add_paragraph()
                if i == 0: target_paragraph.paragraph_format.space_before = Pt(SPACE_ABOVE_FIRST_IMAGE_PT)
                if CENTER_IMAGES: target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                target_run = target_paragraph.add_run()

            try:
                fotos_processadas += 1
                if total_fotos_a_processar > 0:
                    progresso_percent = int((fotos_processadas / total_fotos_a_processar) * 100)
                    progress_callback(progresso_percent)
                    status_callback(f"Lote {idx_lote + 1}/{total_lotes}: Img {i + 1}/{total_imagens_no_lote} ({progresso_percent}%)")
                
                decision_dict = final_crop_decisions.get(caminho_imagem_completo, {})
                crop_decision = decision_dict.get("crop", "skip")
                rotation_decision = decision_dict.get("rotation", 0)
                brightness_decision = decision_dict.get("brightness", 1.0)

                with Image.open(caminho_imagem_completo) as pil_image:
                    # 0. Ler Data EXIF (Antes de qualquer manipulação)
                    timestamp_text = ""
                    if ADD_TIMESTAMP and hasattr(pil_image, '_getexif'):
                        exif_data = pil_image._getexif()
                        if exif_data:
                            # Tag 36867 = DateTimeOriginal, 306 = DateTime
                            date_str = exif_data.get(36867) or exif_data.get(306)
                            if date_str:
                                # Formato padrão EXIF: YYYY:MM:DD HH:MM:SS
                                # Vamos formatar para DD/MM/YYYY HH:MM
                                try:
                                    parts = date_str.split(' ')
                                    date_parts = parts[0].split(':')
                                    time_parts = parts[1].split(':')
                                    timestamp_text = f"{date_parts[2]}/{date_parts[1]}/{date_parts[0]} {time_parts[0]}:{time_parts[1]}"
                                except:
                                    timestamp_text = date_str # Fallback se falhar o parse

                    # 1. Rotação
                    if rotation_decision == 90: pil_image = pil_image.transpose(Image.ROTATE_270)
                    elif rotation_decision == 180: pil_image = pil_image.transpose(Image.ROTATE_180)
                    elif rotation_decision == 270: pil_image = pil_image.transpose(Image.ROTATE_90)
                    
                    # 2. Brilho
                    if brightness_decision != 1.0:
                        enhancer = ImageEnhance.Brightness(pil_image)
                        pil_image = enhancer.enhance(brightness_decision)
                    
                    # 3. Corte
                    nw, nh = pil_image.size
                    if crop_decision != "skip" and isinstance(crop_decision, dict):
                        rel = crop_decision
                        x1, y1 = round(rel['x1']*nw), round(rel['y1']*nh)
                        x2, y2 = round(rel['x2']*nw), round(rel['y2']*nh)
                        if x1 < x2 and y1 < y2:
                            pil_image = pil_image.crop((x1, y1, x2, y2))
                    
                    # --- APLICAÇÃO DO CARIMBO (TIMESTAMP) ---
                    if timestamp_text:
                        draw = ImageDraw.Draw(pil_image)
                        w, h = pil_image.size
                        
                        # Calcula tamanho da fonte dinâmico (3% da altura da imagem)
                        font_size = int(h * 0.03)
                        if font_size < 10: font_size = 10
                        
                        try:
                            # Tenta usar Arial, senão usa padrão
                            font = ImageFont.truetype("arial.ttf", font_size)
                        except IOError:
                            font = ImageFont.load_default()

                        # Calcula posição (Canto inferior direito com margem)
                        # bbox = left, top, right, bottom
                        bbox = draw.textbbox((0, 0), timestamp_text, font=font)
                        text_w = bbox[2] - bbox[0]
                        text_h = bbox[3] - bbox[1]
                        
                        x = w - text_w - (font_size * 0.5) # Margem de 0.5x o tamanho da fonte
                        y = h - text_h - (font_size * 0.5)
                        
                        # Desenha contorno preto (stroke) para legibilidade
                        stroke_width = max(1, int(font_size / 15))
                        draw.text((x, y), timestamp_text, font=font, fill="yellow", stroke_width=stroke_width, stroke_fill="black")
                    # -----------------------------------------

                    # 4. Convert to CV2
                    if pil_image.mode == 'RGBA': pil_image = pil_image.convert('RGB')
                    final_image_data = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
                
                if final_image_data is None or final_image_data.size == 0:
                    raise IOError("Imagem vazia.")

                # Encode e Inserção (Igual a antes)
                is_success, buffer = cv2.imencode(".jpg", final_image_data, [cv2.IMWRITE_JPEG_QUALITY, 90])
                image_stream = io.BytesIO(buffer)
                
                try:
                    with Image.open(caminho_imagem_completo) as img:
                        dpi = img.info.get('dpi', (96, 96))[0]
                except: dpi = 96
                
                fh, fw, _ = final_image_data.shape
                aspect = fh / fw
                nat_w_cm = (fw / dpi) * 2.54 if dpi > 0 else fw / 37.8

                max_w = limit_width_cm
                max_h = max_w * aspect
                if max_h > limit_height_cm:
                    max_h = limit_height_cm
                    max_w = max_h / aspect
                
                if FORCE_RESIZE_SMALLER_IMAGES:
                    w_cm, h_cm = max_w, max_h
                else:
                    w_cm = min(nat_w_cm, max_w)
                    h_cm = w_cm * aspect
                
                run.add_picture(image_stream, width=Cm(w_cm), height=Cm(h_cm))
                
                if not SIDE_BY_SIDE_LAYOUT:
                    cursor.addnext(target_paragraph._element)
                    cursor = target_paragraph._element

            except Exception as e:
                status_callback(f"AVISO: Erro na img {os.path.basename(caminho_imagem_completo)}: {e}")
                try: 
                    if not SIDE_BY_SIDE_LAYOUT: 
                        target_paragraph._element.getparent().remove(target_paragraph._element)
                except: pass
    
    try:
        status_callback("Salvando documento final...")
        doc.save(caminho_documento_final)
        return "Processo concluído com sucesso!"
    except Exception as e:
        raise Exception(f"ERRO CRÍTICO ao salvar: {e}")